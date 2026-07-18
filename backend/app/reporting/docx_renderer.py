"""Minimal DOCX renderer restricted to explicitly synthetic templates."""

from pathlib import Path

from docx import Document

from app.reporting.models import ReportDraft
from app.reporting.templates import TemplateRegistry
from app.reporting.validation import PlaceholderValidator


class ReportRenderingError(RuntimeError):
    pass


class SyntheticDOCXRenderer:
    def __init__(self, registry: TemplateRegistry) -> None:
        self.registry = registry

    def render(self, draft: ReportDraft, replacements: dict[str, str], output_path: Path) -> Path:
        template = self.registry.get(draft.request.template_id)
        if template is None:
            raise ReportRenderingError("report template was not registered")
        if not template.synthetic_test_only:
            raise ReportRenderingError("official template compatibility is not verified")
        validation = PlaceholderValidator().validate(template)
        if not validation.valid:
            raise ReportRenderingError(validation.issues[0].message)
        document = Document(template.path)
        if "SYNTHETIC_TEST_DATA" not in self._all_text(document):
            raise ReportRenderingError("synthetic template marker is missing")
        self._replace(document, replacements)
        for section in draft.sections:
            document.add_heading(section.heading, level=1)
            document.add_paragraph(section.content)
        for report_table in draft.tables:
            if report_table.title:
                document.add_paragraph(report_table.title)
            table = document.add_table(rows=1, cols=len(report_table.headers))
            for index, header in enumerate(report_table.headers):
                table.rows[0].cells[index].text = header
            for values in report_table.rows:
                cells = table.add_row().cells
                for index, value in enumerate(values[: len(cells)]):
                    cells[index].text = value
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path

    @staticmethod
    def _all_text(document: Document) -> str:
        return "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        )

    @staticmethod
    def _replace(document: Document, replacements: dict[str, str]) -> None:
        paragraphs = list(document.paragraphs)
        paragraphs.extend(
            cell.paragraphs for table in document.tables for row in table.rows for cell in row.cells
        )
        flattened = []
        for item in paragraphs:
            flattened.extend(item if isinstance(item, list) else [item])
        for paragraph in flattened:
            text = paragraph.text
            for name, value in replacements.items():
                text = text.replace(f"{{{{{name}}}}}", value)
                text = text.replace(f"{{{{ {name} }}}}", value)
            if text != paragraph.text:
                paragraph.text = text
