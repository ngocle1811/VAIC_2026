"""Create traceable DOCX drafts without claiming official-template compatibility."""

from decimal import Decimal, InvalidOperation
from pathlib import Path
from uuid import uuid4

from docx import Document

from app.models.operational_report import OperationalReportRecord
from app.operational_data.models import DataClassification
from app.reporting.models import ReportStatus
from app.reporting.repository import GeneratedReportRepository


class ReportGenerationService:
    def __init__(self, repository: GeneratedReportRepository, output_dir: Path) -> None:
        self.repository = repository
        self.output_dir = output_dir

    def generate(self, source: OperationalReportRecord, template_id: str):
        if source.classification != DataClassification.SYNTHETIC_TEST_DATA.value:
            raise ValueError("official report templates and mappings are not approved")
        report_id = str(uuid4())
        path = self.output_dir / source.domain / f"{report_id}.docx"
        path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.add_heading("BÁO CÁO NHÁP – CẦN KIỂM TRA", level=0)
        document.add_paragraph("SYNTHETIC_TEST_DATA · NOT_OFFICIAL")
        document.add_paragraph(f"Lĩnh vực: {source.domain}")
        document.add_paragraph(f"Kỳ báo cáo: {source.period_label or source.period_start}")
        document.add_paragraph(f"Đơn vị: {source.organization_name}")
        document.add_paragraph(f"Nguồn dữ liệu vận hành: {source.id}")
        document.add_heading("Số liệu đã chuẩn hóa", level=1)
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Trường dữ liệu"
        table.rows[0].cells[1].text = "Giá trị nguồn"
        for name, value in source.values.items():
            cells = table.add_row().cells
            cells[0].text = name
            cells[1].text = "" if value is None else str(value)
        if source.records:
            document.add_heading("Danh sách bản ghi", level=1)
            headers = list(dict.fromkeys(key for row in source.records for key in row))
            task_table = document.add_table(rows=1, cols=len(headers))
            for index, header in enumerate(headers):
                task_table.rows[0].cells[index].text = header
            for row in source.records:
                cells = task_table.add_row().cells
                for index, header in enumerate(headers):
                    cells[index].text = "" if row.get(header) is None else str(row.get(header))
        document.add_heading("Kết quả kiểm tra", level=1)
        if source.issues:
            for issue in source.issues:
                document.add_paragraph(
                    f"[{issue.get('severity', 'warning')}] {issue.get('code')}: "
                    f"{issue.get('message')}"
                )
        else:
            document.add_paragraph("Không phát hiện lỗi theo bộ quy tắc thử nghiệm đã mã hóa.")
        document.add_paragraph(
            "Báo cáo này chưa tương thích được xác nhận với biểu mẫu chính thức "
            "và phải được người có thẩm quyền kiểm tra."
        )
        document.save(path)
        validation = {
            "numeric_sources_valid": self._validate_numeric_cells(source.values),
            "source_id_present": True,
            "official_template_compatibility_verified": False,
            "human_review_required": True,
        }
        return self.repository.create(
            id=report_id,
            operational_report_id=source.id,
            domain=source.domain,
            template_id=template_id,
            status=ReportStatus.NEEDS_REVIEW.value,
            artifact_path=str(path),
            validation_result=validation,
        )

    @staticmethod
    def _validate_numeric_cells(values: dict) -> bool:
        for value in values.values():
            if value is None or isinstance(value, (str, bool)):
                continue
            try:
                Decimal(str(value))
            except (InvalidOperation, ValueError):
                return False
        return True
