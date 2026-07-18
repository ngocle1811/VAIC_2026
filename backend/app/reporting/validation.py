"""Deterministic placeholder, numeric-source, and citation validation."""

import re
from decimal import Decimal, InvalidOperation

from docx import Document

from app.operational_data.models import OperationalReport
from app.reporting.models import (
    ReportDraft,
    ReportValidationIssue,
    ReportValidationResult,
)
from app.reporting.templates import ReportTemplate

_PLACEHOLDER = re.compile(r"\{\{\s*([A-Za-z0-9_.-]+)\s*\}\}")


def _document_text(document: Document) -> str:
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


class PlaceholderValidator:
    def validate(self, template: ReportTemplate) -> ReportValidationResult:
        document = Document(template.path)
        found = set(_PLACEHOLDER.findall(_document_text(document)))
        missing = sorted(template.required_placeholders - found)
        issues = [
            ReportValidationIssue(
                code="missing_placeholder",
                message=f"Required placeholder is missing: {name}",
                location=name,
            )
            for name in missing
        ]
        return ReportValidationResult(valid=not issues, issues=issues)


class ReportEvidenceValidator:
    def validate(
        self, draft: ReportDraft, operational_report: OperationalReport
    ) -> ReportValidationResult:
        issues = []
        allowed_sources = set(draft.request.allowed_source_ids)
        for section_index, section in enumerate(draft.sections):
            for citation in section.citations:
                if citation.source_id not in allowed_sources:
                    issues.append(
                        ReportValidationIssue(
                            code="unsupported_citation",
                            message=f"Source was not retrieved: {citation.source_id}",
                            location=f"sections[{section_index}]",
                        )
                    )
            issues.extend(
                self._numeric_issues(
                    section.numeric_sources,
                    operational_report,
                    f"sections[{section_index}]",
                )
            )
        for table_index, table in enumerate(draft.tables):
            issues.extend(
                self._numeric_issues(
                    table.numeric_sources,
                    operational_report,
                    f"tables[{table_index}]",
                )
            )
        if not draft.request.human_review_required:
            issues.append(
                ReportValidationIssue(
                    code="human_review_required",
                    message="Report publication requires human review.",
                )
            )
        return ReportValidationResult(valid=not issues, issues=issues)

    @staticmethod
    def _numeric_issues(references, operational_report, location):
        issues = []
        for reference in references:
            source = operational_report.values.get(reference.field_name)
            try:
                matches = source is not None and Decimal(str(source)) == reference.value
            except (InvalidOperation, ValueError):
                matches = False
            if not matches:
                issues.append(
                    ReportValidationIssue(
                        code="numeric_source_mismatch",
                        message=f"Value does not match operational field: {reference.field_name}",
                        location=location,
                    )
                )
        return issues
