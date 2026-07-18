"""Deterministic extractors for the repository's explicitly synthetic reports."""

from __future__ import annotations

import calendar
import hashlib
import re
import unicodedata
from datetime import date, datetime
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any
from uuid import NAMESPACE_URL, uuid5

import fitz
from docx import Document
from openpyxl import load_workbook

from app.operational_data.models import (
    CommonReportMetadata,
    DataClassification,
    IssueSeverity,
    OperationalDomain,
    OperationalProcessingStatus,
    OperationalReport,
    OrganizationMetadata,
    ReportingPeriod,
    ReportingPeriodKind,
    SourceProvenance,
    ValidationIssue,
)


class OperationalExtractionError(ValueError):
    """Raised when a supported file cannot be deterministically interpreted."""


def _key(value: object) -> str:
    text = unicodedata.normalize("NFC", str(value or "")).strip().casefold()
    return re.sub(r"\s+", " ", text).rstrip(":")


FIELD_ALIASES: dict[OperationalDomain, dict[str, str]] = {
    OperationalDomain.POPULATION: {
        _key(alias): field
        for field, aliases in {
            "population_start": ["Dân số đầu kỳ"],
            "arrivals": ["Chuyển đến trong kỳ"],
            "departures": ["Chuyển đi trong kỳ"],
            "births": ["Số trẻ sinh"],
            "deaths": ["Số người chết"],
            "population_total": ["Dân số cuối kỳ", "Tổng số nhân khẩu", "Tổng dân số"],
            "households": ["Số hộ"],
            "male": ["Nam"],
            "female": ["Nữ"],
            "permanent": ["Thường trú"],
            "temporary": ["Tạm trú", "Số người tạm trú"],
            "temporarily_absent": ["Tạm vắng"],
        }.items()
        for alias in aliases
    },
    OperationalDomain.COMPLAINTS: {
        _key(alias): field
        for field, aliases in {
            "beginning_cases": ["Đơn tồn đầu kỳ"],
            "received_cases": [
                "Đơn tiếp nhận trong kỳ",
                "Tổng số đơn tiếp nhận",
                "Số đơn tiếp nhận",
            ],
            "complaint_cases": ["Đơn khiếu nại"],
            "denunciation_cases": ["Đơn tố cáo"],
            "petition_cases": ["Đơn kiến nghị/phản ánh"],
            "within_authority": ["Đơn thuộc thẩm quyền"],
            "outside_authority": ["Đơn không thuộc thẩm quyền", "Đơn không thuộc thẩm"],
            "resolved_cases": ["Đơn đã giải quyết"],
            "ending_cases": ["Đơn tồn cuối kỳ"],
            "citizen_receptions": ["Số lượt tiếp công dân"],
            "citizens_received": ["Số người được tiếp"],
            "crowd_cases": ["Vụ việc đông người"],
        }.items()
        for alias in aliases
    },
    OperationalDomain.TASKS: {
        _key(alias): field
        for field, aliases in {
            "task_total": ["Tổng số nhiệm vụ", "Tổng số nhiệm vụ theo dõi"],
            "completed_on_time": ["Hoàn thành đúng hạn"],
            "completed_late": ["Hoàn thành quá hạn"],
            "in_progress": ["Đang thực hiện"],
            "not_started": ["Chưa thực hiện"],
            "overdue": ["Quá hạn chưa hoàn thành"],
            "completion_rate": ["Tỷ lệ hoàn thành"],
            "average_progress": ["Tiến độ bình quân"],
        }.items()
        for alias in aliases
    },
}

TASK_ALIASES = {
    _key(alias): field
    for field, aliases in {
        "task_id": ["Mã nhiệm vụ", "Mã"],
        "task_description": ["Nội dung nhiệm vụ"],
        "lead_unit": ["Đơn vị chủ trì"],
        "coordinating_units": ["Đơn vị phối hợp"],
        "assignment_date": ["Ngày giao"],
        "due_date": ["Hạn hoàn thành"],
        "status": ["Trạng thái"],
        "progress_percent": ["Tiến độ (%)", "Tiến độ"],
        "completion_date": ["Ngày hoàn thành"],
        "completion_result": ["Kết quả", "Kết quả/Ghi chú"],
    }.items()
    for alias in aliases
}


def parse_scalar(value: object) -> str | int | Decimal | bool | date | None:
    if value is None or value == "":
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, (int, Decimal, bool, date)):
        return value
    if isinstance(value, float):
        return int(value) if value.is_integer() else Decimal(str(value))
    text = str(value).strip()
    local_date = re.match(r"^(\d{1,2})/(\d{1,2})/(\d{4})$", text)
    if local_date:
        day, month, year = [int(item) for item in local_date.groups()]
        return date(year, month, day)
    iso_date = re.match(r"^(\d{4})-(\d{1,2})-(\d{1,2})$", text)
    if iso_date:
        return date(*[int(item) for item in iso_date.groups()])
    numeric = re.sub(r"\s*(?:người|hộ|đơn|lượt|vụ|nhiệm vụ|%)\s*$", "", text, flags=re.I)
    if re.fullmatch(r"[-+]?\d{1,3}(?:\.\d{3})+", numeric):
        numeric = numeric.replace(".", "")
    elif re.fullmatch(r"[-+]?\d+,\d+", numeric):
        numeric = numeric.replace(",", ".")
    try:
        number = Decimal(numeric)
        return int(number) if number == number.to_integral() else number
    except InvalidOperation:
        return text


class RepositoryOperationalExtractor:
    """Extract tables/text without LLMs and label unapproved mappings for review."""

    def extract(self, source_path: Path, domain: OperationalDomain) -> OperationalReport:
        suffix = source_path.suffix.lower()
        if suffix == ".xlsx":
            rows, text, locations = self._xlsx(source_path)
        elif suffix == ".docx":
            rows, text, locations = self._docx(source_path)
        elif suffix == ".pdf":
            rows, text, locations = self._pdf(source_path)
        else:
            raise OperationalExtractionError(f"unsupported operational file type: {suffix}")
        synthetic = "DỮ LIỆU GIẢ LẬP" in text.upper() or "SYNTHETIC_TEST_DATA" in text.upper()
        classification = (
            DataClassification.SYNTHETIC_TEST_DATA
            if synthetic
            else DataClassification.OFFICIAL_CANDIDATE
        )
        values, records, provenance = self._map_rows(rows, locations, domain, source_path.name)
        period = self._period(text, source_path.name)
        organization = self._organization(text)
        checksum = hashlib.sha256(source_path.read_bytes()).hexdigest()
        report_id = str(uuid5(NAMESPACE_URL, f"{domain.value}:{checksum}"))
        issues: list[ValidationIssue] = []
        if not synthetic:
            issues.append(
                ValidationIssue(
                    code="UNAPPROVED_OFFICIAL_MAPPING",
                    message=(
                        "Official schema mapping is not approved; extracted values "
                        "require human review."
                    ),
                    severity=IssueSeverity.WARNING,
                )
            )
        return OperationalReport(
            metadata=CommonReportMetadata(
                report_id=report_id,
                domain=domain,
                reporting_period=period,
                organization=organization,
                classification=classification,
                schema_version="repository-synthetic-v1" if synthetic else None,
            ),
            values=values,
            records=records,
            provenance=provenance,
            processing_status=OperationalProcessingStatus.EXTRACTED,
            issues=issues,
        )

    def _map_rows(self, rows, locations, domain, filename):
        aliases = FIELD_ALIASES[domain]
        values: dict[str, Any] = {}
        provenance: dict[str, list[SourceProvenance]] = {}
        records: list[dict[str, Any]] = []
        for row_index, row in enumerate(rows):
            normalized = [_key(item) for item in row]
            task_headers = {TASK_ALIASES[item] for item in normalized if item in TASK_ALIASES}
            if domain is OperationalDomain.TASKS and {
                "task_id",
                "task_description",
            }.issubset(task_headers):
                headers = [TASK_ALIASES.get(item) for item in normalized]
                for data_row in rows[row_index + 1 :]:
                    record = {
                        field: parse_scalar(data_row[index] if index < len(data_row) else None)
                        for index, field in enumerate(headers)
                        if field
                    }
                    if re.match(r"^NV-", str(record.get("task_id") or ""), re.I) and record.get(
                        "task_description"
                    ):
                        records.append(record)
                    elif records:
                        break
                break
            if len(row) < 2:
                continue
            field = aliases.get(normalized[0])
            if not field:
                continue
            values[field] = parse_scalar(row[1])
            location = locations.get(row_index, {})
            provenance[field] = [
                SourceProvenance(
                    source_file_id=filename,
                    extraction_method=location.get("method", "deterministic_table"),
                    page_number=location.get("page"),
                    sheet_name=location.get("sheet"),
                    row_number=location.get("row"),
                    evidence=str(row[0]),
                )
            ]
        if not values and not records:
            raise OperationalExtractionError("no recognized fields were found")
        return values, records, provenance

    @staticmethod
    def _xlsx(path: Path):
        workbook = load_workbook(path, data_only=True, read_only=True)
        rows, locations, text = [], {}, []
        for sheet in workbook.worksheets:
            for row_number, values in enumerate(sheet.iter_rows(values_only=True), start=1):
                row = list(values)
                rows.append(row)
                locations[len(rows) - 1] = {
                    "method": "openpyxl",
                    "sheet": sheet.title,
                    "row": row_number,
                }
                text.extend(str(item) for item in row if item is not None)
        workbook.close()
        return rows, "\n".join(text), locations

    @staticmethod
    def _docx(path: Path):
        document = Document(path)
        rows, locations = [], {}
        text = [paragraph.text for paragraph in document.paragraphs]
        for table_index, table in enumerate(document.tables, start=1):
            for row_number, row in enumerate(table.rows, start=1):
                values = [cell.text.strip() for cell in row.cells]
                rows.append(values)
                locations[len(rows) - 1] = {
                    "method": "python-docx",
                    "sheet": f"table_{table_index}",
                    "row": row_number,
                }
                text.extend(values)
        return rows, "\n".join(text), locations

    @staticmethod
    def _pdf(path: Path):
        rows, locations, all_lines = [], {}, []
        with fitz.open(path) as document:
            for page_index, page in enumerate(document, start=1):
                lines = [
                    line.strip() for line in page.get_text("text").splitlines() if line.strip()
                ]
                all_lines.extend(lines)
                for index, line in enumerate(lines[:-1]):
                    following = lines[index + 1 : index + 4]
                    value = next(
                        (item for item in following if re.fullmatch(r"[-+]?\d+(?:[.,]\d+)?", item)),
                        lines[index + 1],
                    )
                    rows.append([line, value])
                    locations[len(rows) - 1] = {
                        "method": "PyMuPDF_text",
                        "page": page_index,
                    }
        return rows, "\n".join(all_lines), locations

    @staticmethod
    def _period(text: str, filename: str) -> ReportingPeriod:
        match = re.search(r"thang[_\s-]*(\d{1,2})[_-](\d{4})", filename, re.I)
        if not match:
            match = re.search(r"(?:tháng|kỳ báo cáo)\s*:?[ _-]*(\d{1,2})[/_-](\d{4})", text, re.I)
            if not match:
                raise OperationalExtractionError("reporting month could not be determined")
            month, year = int(match.group(1)), int(match.group(2))
        else:
            month, year = int(match.group(1)), int(match.group(2))
        last_day = calendar.monthrange(year, month)[1]
        return ReportingPeriod(
            kind=ReportingPeriodKind.MONTH,
            start=date(year, month, 1),
            end=date(year, month, last_day),
            label=f"Tháng {month:02d}/{year}",
        )

    @staticmethod
    def _organization(text: str) -> OrganizationMetadata:
        match = re.search(r"(?:Đơn vị báo cáo|Cơ quan báo cáo)\s*:?\s*([^\n]+)", text, re.I)
        name = match.group(1).strip() if match else "Đơn vị chưa xác định"
        return OrganizationMetadata(
            organization_id=str(uuid5(NAMESPACE_URL, _key(name))), organization_name=name
        )
