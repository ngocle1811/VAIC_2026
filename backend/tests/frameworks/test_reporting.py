from decimal import Decimal

import pytest
from docx import Document

from app.reporting.docx_renderer import ReportRenderingError, SyntheticDOCXRenderer
from app.reporting.models import (
    NumericSourceReference,
    ReportCitation,
    ReportDraft,
    ReportRequest,
    ReportSection,
    ReportTable,
)
from app.reporting.templates import InMemoryTemplateRegistry, ReportTemplate
from app.reporting.validation import PlaceholderValidator, ReportEvidenceValidator


def _draft() -> ReportDraft:
    return ReportDraft(
        request=ReportRequest(
            report_id="synthetic-draft",
            domain="tasks",
            operational_report_id="synthetic-report",
            template_id="synthetic-template",
            allowed_source_ids=["SOURCE_1"],
        ),
        title="SYNTHETIC_TEST_DATA",
        reporting_period_label="2099-W01",
        sections=[
            ReportSection(
                heading="Synthetic section",
                content="Synthetic narrative",
                citations=[ReportCitation(source_id="SOURCE_1")],
                numeric_sources=[
                    NumericSourceReference(field_name="metric_alpha", value=Decimal("20"))
                ],
            )
        ],
        tables=[ReportTable(headers=["Field", "Value"], rows=[["alpha", "20"]])],
    )


def test_report_evidence_rejects_numeric_and_citation_mismatches(synthetic_report) -> None:
    draft = _draft()
    assert ReportEvidenceValidator().validate(draft, synthetic_report).valid
    invalid = draft.model_copy(deep=True)
    invalid.sections[0].citations[0].source_id = "SOURCE_99"
    invalid.sections[0].numeric_sources[0].value = Decimal("999")
    result = ReportEvidenceValidator().validate(invalid, synthetic_report)
    assert {issue.code for issue in result.issues} == {
        "unsupported_citation",
        "numeric_source_mismatch",
    }


def test_synthetic_docx_placeholder_validation_and_rendering(tmp_path) -> None:
    source = tmp_path / "synthetic-template.docx"
    document = Document()
    document.add_paragraph("SYNTHETIC_TEST_DATA")
    document.add_paragraph("{{title}} - {{period}}")
    document.save(source)
    template = ReportTemplate(
        template_id="synthetic-template",
        path=source,
        required_placeholders={"title", "period"},
        synthetic_test_only=True,
    )
    assert PlaceholderValidator().validate(template).valid
    output = tmp_path / "rendered.docx"
    SyntheticDOCXRenderer(InMemoryTemplateRegistry([template])).render(
        _draft(), {"title": "Synthetic report", "period": "2099-W01"}, output
    )
    rendered = Document(output)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "Synthetic report - 2099-W01" in text
    assert "Synthetic narrative" in text

    unsafe = template.model_copy(update={"synthetic_test_only": False})
    with pytest.raises(ReportRenderingError, match="not verified"):
        SyntheticDOCXRenderer(InMemoryTemplateRegistry([unsafe])).render(
            _draft(), {}, tmp_path / "unsafe.docx"
        )
