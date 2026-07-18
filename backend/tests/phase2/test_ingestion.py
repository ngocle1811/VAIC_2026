from pathlib import Path

import pytest

from app.models.knowledge_document import ProcessingStatus
from app.rag.embeddings.service import EmbeddingService
from app.rag.ingestion.pipeline import IngestionError, KnowledgeBaseIngestionPipeline
from app.rag.models import DocumentDomain, DocumentType
from app.rag.parsers.factory import ParserFactory
from app.rag.pipeline import RAGPhase1Pipeline
from app.repositories.knowledge_document_repository import KnowledgeDocumentRepository
from app.schemas.knowledge_document import IngestionOutcome, IngestionRequest
from app.services.knowledge_base_storage import KnowledgeBaseStorage


def _docx(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading("QUY TRÌNH BÁO CÁO", level=1)
    document.add_paragraph("Nội dung báo cáo định kỳ và trách nhiệm thực hiện.")
    document.save(path)


def _pipeline(db_session, tmp_path, embedding_service, vector_store, phase1=None):
    return KnowledgeBaseIngestionPipeline(
        repository=KnowledgeDocumentRepository(db_session),
        storage=KnowledgeBaseStorage(tmp_path / "managed"),
        phase1_pipeline=phase1
        or RAGPhase1Pipeline(parser_factory=ParserFactory(prefer_docling=False)),
        embedding_service=embedding_service,
        vector_store=vector_store,
    )


def _request(path: Path, force: bool = False) -> IngestionRequest:
    return IngestionRequest(
        source_path=path,
        document_name="Quy trình",
        document_type=DocumentType.GUIDELINE,
        domain=DocumentDomain.COMMON,
        force_reindex=force,
    )


def test_success_duplicate_and_force_reindex(
    db_session, tmp_path: Path, embedding_service, fake_provider, vector_store
) -> None:
    source = tmp_path / "source.docx"
    _docx(source)
    pipeline = _pipeline(db_session, tmp_path, embedding_service, vector_store)
    first = pipeline.ingest(_request(source))
    assert first.outcome == IngestionOutcome.INDEXED
    assert first.processing_status == ProcessingStatus.READY
    assert vector_store.count_by_document_id(first.document_id) == first.chunk_count
    calls = fake_provider.document_calls
    duplicate = pipeline.ingest(_request(source))
    assert duplicate.outcome == IngestionOutcome.SKIPPED_DUPLICATE
    assert fake_provider.document_calls == calls
    reindexed = pipeline.ingest(_request(source, force=True))
    assert reindexed.outcome == IngestionOutcome.REINDEXED
    assert vector_store.count_by_document_id(first.document_id) == first.chunk_count


def test_successful_pdf_ingestion(
    db_session, tmp_path: Path, embedding_service, vector_store
) -> None:
    import fitz

    source = tmp_path / "legal.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Dieu 1. Noi dung bao cao")
    pdf.save(source)
    pdf.close()
    pipeline = _pipeline(db_session, tmp_path, embedding_service, vector_store)
    request = IngestionRequest(
        source_path=source,
        document_name="Van ban",
        document_type=DocumentType.LEGAL,
        domain=DocumentDomain.COMMON,
    )
    result = pipeline.ingest(request)
    assert result.processing_status == ProcessingStatus.READY
    assert result.chunk_count > 0


class EmptyPhase1:
    def process(self, *args, **kwargs):
        return []


class FailingProvider:
    model_name = "Vietnamese_Embedding"

    def embed_documents(self, texts):
        raise RuntimeError("provider failure without document text")

    def embed_query(self, query):
        raise RuntimeError("provider failure")


class MismatchVectorStore:
    def collection_exists(self):
        return True

    def create_collection_if_missing(self, vector_size):
        pass

    def upsert_chunks(self, chunks):
        pass

    def search(self, query_vector, filters, top_k):
        return []

    def delete_by_document_id(self, document_id):
        pass

    def count_by_document_id(self, document_id):
        return 0


@pytest.mark.parametrize("failure", ["empty", "embedding", "count"])
def test_ingestion_failures_mark_document_failed(
    failure, db_session, tmp_path: Path, embedding_service, vector_store
) -> None:
    source = tmp_path / f"{failure}.docx"
    _docx(source)
    phase1 = EmptyPhase1() if failure == "empty" else None
    service = (
        EmbeddingService(FailingProvider(), normalize=False)
        if failure == "embedding"
        else embedding_service
    )
    store = MismatchVectorStore() if failure == "count" else vector_store
    pipeline = _pipeline(db_session, tmp_path, service, store, phase1)
    with pytest.raises(IngestionError):
        pipeline.ingest(_request(source))
    document = KnowledgeDocumentRepository(db_session).list_documents()[0]
    assert document.processing_status == ProcessingStatus.FAILED
    expected_stage = {
        "empty": "chunk_validation",
        "embedding": "embedding",
        "count": "vector_verification",
    }[failure]
    assert document.failed_stage == expected_stage
